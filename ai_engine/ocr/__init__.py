"""
OCR Module for extracting text from documents.
Supports PDF, scanned images, and various image formats.
Languages: Uzbek (Latin/Cyrillic), Russian
"""

import os
import io
import unicodedata
import logging
from typing import Optional, Tuple, List
from pathlib import Path

import pytesseract
from PIL import Image, ImageEnhance, ImageFilter
from pdf2image import convert_from_path, convert_from_bytes
import pdfplumber
try:
    import fitz  # PyMuPDF, optional fallback renderer
except ImportError:  # pragma: no cover
    fitz = None

logger = logging.getLogger(__name__)


class OCRProcessor:
    """
    OCR processor for extracting text from documents.
    Uses Tesseract for OCR and pdfplumber for native PDF text.
    """
    # Default tessdata path
    DEFAULT_TESSDATA = '/usr/share/tesseract-ocr/5/tessdata'
    
    # Language codes for Tesseract
    LANGUAGE_MAP = {
        'uz-latn': 'uzb_latn',  # Uzbek Latin
        'uz-cyrl': 'uzb',       # Uzbek Cyrillic
        'ru': 'rus',            # Russian
        'eng': 'eng',           # English
    }
    
    def __init__(self, languages: List[str] = None):
        """
        Initialize OCR processor.
        
        Args:
            languages: List of language codes to use for OCR
        """
        os.environ.setdefault('TESSDATA_PREFIX', self.DEFAULT_TESSDATA)
        self.languages = self._build_language_list(languages or ['uzb', 'uzb_latn', 'rus'])
        self.tesseract_lang = '+'.join(self.languages)
        # Env-driven configs
        try:
            self.max_pages = int(os.environ.get('OCR_PAGES_MAX', '0'))
        except Exception:
            self.max_pages = 0
        try:
            self.ocr_dpi = int(os.environ.get('OCR_DPI', '400'))
        except Exception:
            self.ocr_dpi = 400
        # Optional PaddleOCR fallback
        self.use_paddle = os.environ.get('USE_PADDLE', '0') in ('1', 'true', 'True')
        self._paddle = None
        if self.use_paddle:
            try:
                self._paddle = PaddleOCRProcessor()
                if not getattr(self._paddle, 'available', False):
                    self._paddle = None
                    logger.warning("USE_PADDLE=1 set but PaddleOCR is not available")
            except Exception as e:
                logger.warning(f"Failed to initialize PaddleOCR: {e}")
        # Optional chunking to limit memory/time on long PDFs
        try:
            self.chunk_size = int(os.environ.get('OCR_CHUNK_SIZE', '0'))
        except Exception:
            self.chunk_size = 0

        # OCR behavior toggles
        self.force_ocr = os.environ.get('OCR_FORCE', '0') in ('1', 'true', 'True')
        try:
            self.min_quality = float(os.environ.get('OCR_QUALITY_MIN', '150'))
        except Exception:
            self.min_quality = 150.0
        try:
            self.max_gibberish = float(os.environ.get('OCR_GIBBERISH_MAX', '0.55'))
        except Exception:
            self.max_gibberish = 0.55
        
        # Verify Tesseract is installed
        try:
            pytesseract.get_tesseract_version()
        except Exception as e:
            logger.warning(f"Tesseract not properly configured: {e}")

        # Simple character sets for quick quality checks
        self._cyrillic_chars = set('абвгдеёжзийклмнопрстуфхцчшщъыьэюяАБВГДЕЁЖЗИЙКЛМНОПРСТУФХЦЧШЩЪЫЬЭЮЯ')
        self._latin_chars = set('abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ')

    def _build_language_list(self, requested: List[str]) -> List[str]:
        """Return only the languages that are installed."""
        available = set()
        try:
            available = set(pytesseract.get_languages(config=''))
        except Exception as e:
            logger.warning(f"Unable to enumerate tesseract languages: {e}")

        normalized = []
        for lang in requested:
            if lang in available:
                normalized.append(lang)
                continue
            mapped = self.LANGUAGE_MAP.get(lang)
            if mapped and mapped in available:
                normalized.append(mapped)
                continue
            if lang in self.LANGUAGE_MAP.values() and lang in available:
                normalized.append(lang)
                continue
            logger.debug(f"Tesseract language '{lang}' is not installed")

        if not normalized:
            fallback = 'eng'
            logger.warning(
                f"None of the requested OCR languages are available, falling back to '{fallback}'"
            )
            normalized = [fallback]

        return normalized
    
    def extract_text_from_file(self, file_path: str) -> Tuple[str, float, bool]:
        """
        Extract text from a file (PDF, image, or Word document).
        
        Args:
            file_path: Path to the file
            
        Returns:
            Tuple of (extracted_text, confidence_score, is_scanned)
        """
        file_path = Path(file_path)
        extension = file_path.suffix.lower()
        
        if extension == '.pdf':
            return self._process_pdf(file_path)
        elif extension in ['.jpg', '.jpeg', '.png', '.tiff', '.bmp', '.gif']:
            return self._process_image(file_path)
        elif extension in ['.docx', '.doc']:
            return self._process_word(file_path)
        else:
            raise ValueError(f"Unsupported file format: {extension}")
    
    def extract_text_from_bytes(self, file_bytes: bytes, file_type: str) -> Tuple[str, float, bool]:
        """
        Extract text from file bytes.
        
        Args:
            file_bytes: File content as bytes
            file_type: File extension (e.g., 'pdf', 'jpg')
            
        Returns:
            Tuple of (extracted_text, confidence_score, is_scanned)
        """
        if file_type == 'pdf':
            return self._process_pdf_bytes(file_bytes)
        elif file_type in ['jpg', 'jpeg', 'png', 'tiff', 'bmp']:
            return self._process_image_bytes(file_bytes)
        else:
            raise ValueError(f"Unsupported file format: {file_type}")
    
    def _process_pdf(self, file_path: Path) -> Tuple[str, float, bool]:
        """Process PDF file."""
        text = ""
        is_scanned = False
        confidence = 1.0
        
        # First try to extract native text
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            logger.warning(f"pdfplumber extraction failed: {e}")

        # Normalize to reduce ligatures and odd encodings, then score
        native_text = self._normalize_text(text.strip())
        native_score = self._text_quality_score(native_text)
        gibberish = self._gibberish_ratio(native_text)

        # Decide if OCR fallback is needed even when some text exists
        need_ocr = self.force_ocr or self._needs_ocr(native_text, native_score, gibberish)

        if need_ocr:
            ocr_text, ocr_conf = self._ocr_pdf(file_path)
            ocr_text = ocr_text.strip()
            ocr_text = self._normalize_text(ocr_text)
            ocr_score = self._text_quality_score(ocr_text)

            # Pick the better of native vs OCR based on quality score
            if ocr_score > native_score * 1.15 or not native_text:
                return ocr_text, ocr_conf, True
            # Keep native but still return best confidence indicator
            confidence = max(confidence, ocr_conf)

        return native_text, confidence, False
    
    def _process_pdf_bytes(self, file_bytes: bytes) -> Tuple[str, float, bool]:
        """Process PDF from bytes."""
        text = ""
        is_scanned = False
        confidence = 1.0
        
        # First try to extract native text
        try:
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            logger.warning(f"pdfplumber extraction failed: {e}")

        native_text = self._normalize_text(text.strip())
        native_score = self._text_quality_score(native_text)
        gibberish = self._gibberish_ratio(native_text)
        need_ocr = self.force_ocr or self._needs_ocr(native_text, native_score, gibberish)

        if need_ocr:
            # Chunked path if enabled and PyMuPDF available
            if self.chunk_size and fitz:
                doc = fitz.open(stream=file_bytes, filetype='pdf')
                total_pages = len(doc)
                last_page_limit = self.max_pages if self.max_pages > 0 else total_pages
                last_page_limit = min(last_page_limit, total_pages)
                texts: List[str] = []
                confs: List[float] = []
                start = 1
                while start <= last_page_limit:
                    end = min(start + self.chunk_size - 1, last_page_limit)
                    images: List[Image.Image] = []
                    try:
                        images = convert_from_bytes(
                            file_bytes, dpi=self.ocr_dpi, first_page=start, last_page=end
                        )
                    except Exception as e:
                        logger.warning(f"pdf2image bytes chunk {start}-{end} failed: {e}")
                    if not images:
                        for p in range(start - 1, end):
                            page = doc.load_page(p)
                            pix = page.get_pixmap(dpi=self.ocr_dpi)
                            images.append(Image.open(io.BytesIO(pix.tobytes('png'))))
                    if images:
                        t, c = self._ocr_images(images)
                        texts.append(t)
                        confs.append(c)
                    start = end + 1
                ocr_text = "\n".join(texts)
                ocr_conf = sum(confs) / len(confs) if confs else 0.0
            else:
                images: List[Image.Image] = []
                try:
                    images = convert_from_bytes(
                        file_bytes,
                        dpi=self.ocr_dpi,
                        first_page=1,
                        last_page=(self.max_pages if self.max_pages > 0 else None)
                    )
                except Exception as e:
                    logger.warning(f"pdf2image (bytes) failed: {e}")
                if not images and fitz:
                    images = self._render_pdf_with_pymupdf(file_bytes=file_bytes)
                if not images:
                    raise RuntimeError("PDF rendering failed (bytes)")
                ocr_text, ocr_conf = self._ocr_images(images)
            ocr_text = self._normalize_text(ocr_text.strip())
            ocr_score = self._text_quality_score(ocr_text)

            if ocr_score > native_score * 1.15 or not native_text:
                return ocr_text, ocr_conf, True
            confidence = max(confidence, ocr_conf)

        return native_text, confidence, False
    
    def _ocr_pdf(self, file_path: Path) -> Tuple[str, float]:
        """Perform OCR on PDF pages with optional chunking."""
        # Chunked path if enabled and PyMuPDF available
        if self.chunk_size and fitz:
            doc = fitz.open(file_path)
            total_pages = len(doc)
            last_page_limit = self.max_pages if self.max_pages > 0 else total_pages
            last_page_limit = min(last_page_limit, total_pages)
            texts: List[str] = []
            confs: List[float] = []
            start = 1
            while start <= last_page_limit:
                end = min(start + self.chunk_size - 1, last_page_limit)
                images: List[Image.Image] = []
                try:
                    images = convert_from_path(
                        str(file_path), dpi=self.ocr_dpi, first_page=start, last_page=end
                    )
                except Exception as e:
                    logger.warning(f"pdf2image chunk {start}-{end} failed: {e}")
                if not images:
                    for p in range(start - 1, end):
                        page = doc.load_page(p)
                        pix = page.get_pixmap(dpi=self.ocr_dpi)
                        images.append(Image.open(io.BytesIO(pix.tobytes('png'))))
                if images:
                    t, c = self._ocr_images(images)
                    texts.append(t)
                    confs.append(c)
                start = end + 1
            avg_conf = sum(confs) / len(confs) if confs else 0.0
            return "\n".join(texts), avg_conf
        # Non-chunked path
        images: List[Image.Image] = []
        try:
            images = convert_from_path(
                str(file_path),
                dpi=self.ocr_dpi,
                first_page=1,
                last_page=(self.max_pages if self.max_pages > 0 else None)
            )
        except Exception as e:
            logger.warning(f"pdf2image failed: {e}")
        if not images and fitz:
            images = self._render_pdf_with_pymupdf(file_path=file_path)
        if not images:
            raise RuntimeError("PDF rendering failed")
        return self._ocr_images(images)
    
    def _ocr_images(self, images: List[Image.Image]) -> Tuple[str, float]:
        """Perform OCR on a list of images."""
        all_text = []
        confidences = []
        
        for img in images:
            t_text, t_conf = self._ocr_image_best(img)
            best_text, best_conf = t_text, t_conf

            # If Tesseract result is weak and PaddleOCR is enabled, try fallback
            if self._paddle and (best_conf < 0.55 or len(best_text) < 100):
                try:
                    import tempfile
                    with tempfile.NamedTemporaryFile(suffix='.png', delete=True) as tmp:
                        img.save(tmp.name, format='PNG')
                        p_text, p_conf = self._paddle.extract_text(tmp.name)
                    # Choose better by a simple score (confidence, length)
                    def score(txt, conf):
                        return (conf or 0) * 1000 + len(txt)
                    if score(p_text, p_conf) > score(best_text, best_conf):
                        best_text, best_conf = p_text, p_conf
                except Exception as e:
                    logger.debug(f"Paddle fallback failed: {e}")

            all_text.append(best_text)
            if best_conf:
                confidences.append(best_conf * 100 if best_conf <= 1 else best_conf)
        
        avg_confidence = sum(confidences) / len(confidences) if confidences else 0
        return "\n".join(all_text), avg_confidence / 100

    def _ocr_image_best(self, img: Image.Image) -> Tuple[str, float]:
        """Try multiple Tesseract configs (PSM) and pick best by confidence/length."""
        candidates = []
        # Try a few common PSM values: 1=Auto OSD, 3=Auto, 6=Block, 4=Column
        for psm in (1, 3, 6, 4):
            config = f"--psm {psm}"
            try:
                data = pytesseract.image_to_data(
                    img,
                    lang=self.tesseract_lang,
                    config=config,
                    output_type=pytesseract.Output.DICT
                )
                text = pytesseract.image_to_string(img, lang=self.tesseract_lang, config=config)
                conf_values = [int(c) for c in data.get('conf', []) if str(c).isdigit() and int(c) > 0]
                avg_conf = (sum(conf_values) / len(conf_values) / 100) if conf_values else 0
                candidates.append((text, avg_conf, psm))
            except Exception as e:
                logger.debug(f"Tesseract psm={psm} failed: {e}")
        if not candidates:
            return "", 0.0
        # Rank by confidence first, then by text length
        candidates.sort(key=lambda x: (x[1], len(x[0])), reverse=True)
        best_text, best_conf, _ = candidates[0]
        return best_text, best_conf
    
    def _process_image(self, file_path: Path) -> Tuple[str, float, bool]:
        """Process image file."""
        img = Image.open(file_path)
        text, confidence = self._ocr_single_image(img)
        return text, confidence, True
    
    def _process_image_bytes(self, file_bytes: bytes) -> Tuple[str, float, bool]:
        """Process image from bytes."""
        img = Image.open(io.BytesIO(file_bytes))
        text, confidence = self._ocr_single_image(img)
        return text, confidence, True

    def _ocr_single_image(self, img: Image.Image) -> Tuple[str, float]:
        """Perform OCR on a single image with preprocessing and PSM selection."""
        img = self._preprocess_image(img)
        return self._ocr_image_best(img)

    def _preprocess_image(self, img: Image.Image) -> Image.Image:
        """Improved preprocessing for scanned documents."""
        # 1. Grayscale
        img = img.convert('L')
        # 2. Increase contrast more aggressively for scanned docs
        img = ImageEnhance.Contrast(img).enhance(2.5)
        # 3. Sharpen
        img = ImageEnhance.Sharpness(img).enhance(1.5)
        # 4. Adaptive threshold (better for varied lighting)
        img = img.point(lambda x: 0 if x < 140 else 255, '1')
        # 5. Noise removal (median filter)
        img = img.filter(ImageFilter.MedianFilter(size=3))
        # 6. Scale up small images for better recognition
        width, height = img.size
        if width < 1500:
            ratio = 1500 / width
            img = img.resize((int(width * ratio), int(height * ratio)), Image.Resampling.LANCZOS)
        # 7. Convert back to RGB for Tesseract
        img = img.convert('RGB')
        return img

    def _text_quality_score(self, text: str) -> float:
        """Rough quality score to compare native PDF text vs OCR output."""
        if not text:
            return 0.0
        letters = sum(1 for c in text if c.isalpha())
        words = len(text.split())
        cyrillic = sum(1 for c in text if c in self._cyrillic_chars)
        latin = sum(1 for c in text if c in self._latin_chars)
        script_bonus = 0.0
        if letters:
            cyr_ratio = cyrillic / letters
            lat_ratio = latin / letters
            # Encourage clear script dominance to prefer readable text
            script_bonus = max(cyr_ratio, lat_ratio) * 200
        # Penalize excessive symbol density
        symbols = sum(1 for c in text if not (c.isalpha() or c.isdigit() or c.isspace()))
        symbol_penalty = min(symbols * 0.1, 400)
        # Letters carry the most weight; words add a small bonus; subtract symbol penalty
        return max(0.0, letters + (words * 0.3) + script_bonus - symbol_penalty)

    def _needs_ocr(self, text: str, score: float, gibberish_ratio: float) -> bool:
        """Decide if OCR fallback is needed despite native text."""
        if not text:
            return True
        letters = sum(1 for c in text if c.isalpha())
        # Trigger OCR on poor native quality or excessive gibberish
        if letters < 300:
            return True
        if score < self.min_quality:
            return True
        if text.count('\ufffd') > 0:
            return True
        if gibberish_ratio > self.max_gibberish:
            return True
        return False

    def _normalize_text(self, text: str) -> str:
        """Normalize unicode (reduce ligatures) and collapse excessive spaces."""
        if not text:
            return text
        # Normalize ligatures and compatibility forms
        norm = unicodedata.normalize('NFKC', text)
        # Collapse multiple spaces
        norm = ' '.join(norm.split())
        return norm

    def _gibberish_ratio(self, text: str) -> float:
        """Estimate gibberish by counting vowel-less long tokens."""
        if not text:
            return 1.0
        tokens = [t for t in text.split() if len(t) >= 4]
        if not tokens:
            return 0.0
        vowels = set("aeiouyAEIOUYоаеёиыуэюяОАЕЁИЫУЭЮЯўғқҳЎҒҚҲo'" )
        def is_gib(tok: str) -> bool:
            return not any(ch in vowels for ch in tok)
        gib = sum(1 for t in tokens if is_gib(t))
        return gib / max(1, len(tokens))

    def _render_pdf_with_pymupdf(self, file_path: Optional[Path] = None, file_bytes: Optional[bytes] = None) -> List[Image.Image]:
        """Render PDF pages to images using PyMuPDF fallback (avoids poppler dependency)."""
        if not fitz:
            return []
        if not file_path and file_bytes is None:
            return []
        doc = fitz.open(stream=file_bytes, filetype="pdf") if file_bytes is not None else fitz.open(file_path)
        images: List[Image.Image] = []
        for page in doc:
            pix = page.get_pixmap(dpi=300)
            images.append(Image.open(io.BytesIO(pix.tobytes("png"))))
        return images
    
    def _process_word(self, file_path: Path) -> Tuple[str, float, bool]:
        """Process Word document (.docx and .doc)."""
        extension = file_path.suffix.lower()
        
        if extension == '.docx':
            return self._process_docx(file_path)
        elif extension == '.doc':
            return self._process_doc(file_path)
        else:
            raise ValueError(f"Unsupported Word format: {extension}")
    
    def _process_docx(self, file_path: Path) -> Tuple[str, float, bool]:
        """Process .docx Word document."""
        from docx import Document
        
        doc = Document(file_path)
        text = []
        
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text.append(cell.text)
        
        return "\n".join(text), 1.0, False
    
    def _process_doc(self, file_path: Path) -> Tuple[str, float, bool]:
        """Process old .doc Word document using antiword or textract."""
        import subprocess
        
        # Try antiword first
        try:
            result = subprocess.run(
                ['antiword', str(file_path)],
                capture_output=True,
                text=True,
                timeout=30
            )
            if result.returncode == 0:
                return result.stdout.strip(), 1.0, False
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass
        
        # Try catdoc as fallback
        try:
            result = subprocess.run(
                ['catdoc', str(file_path)],
                capture_output=True,
                text=True,
                timeout=30
            )
            if result.returncode == 0:
                return result.stdout.strip(), 1.0, False
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass
        
        # Try libreoffice as last resort
        try:
            import tempfile
            with tempfile.TemporaryDirectory() as tmpdir:
                result = subprocess.run(
                    ['libreoffice', '--headless', '--convert-to', 'txt:Text', '--outdir', tmpdir, str(file_path)],
                    capture_output=True,
                    text=True,
                    timeout=60
                )
                if result.returncode == 0:
                    txt_file = Path(tmpdir) / (file_path.stem + '.txt')
                    if txt_file.exists():
                        return txt_file.read_text(encoding='utf-8'), 1.0, False
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass
        
        # If all methods fail, try reading as binary and extracting text
        try:
            with open(file_path, 'rb') as f:
                content = f.read()
            
            # Extract printable text from binary
            import re
            # Try to find UTF-8 or Latin-1 text sequences
            text_parts = []
            
            # Try different decodings
            for encoding in ['utf-8', 'cp1251', 'cp1252', 'latin-1']:
                try:
                    decoded = content.decode(encoding, errors='ignore')
                    # Extract words (at least 3 consecutive printable chars)
                    words = re.findall(r'[\w\s.,;:!?()-]{3,}', decoded)
                    if words:
                        text_parts.extend(words)
                        break
                except:
                    continue
            
            if text_parts:
                return ' '.join(text_parts)[:50000], 0.5, False  # Limit text and lower confidence
                
        except Exception as e:
            logger.warning(f"Failed to extract text from .doc file: {e}")
        
        raise ValueError(
            f"Cannot process .doc file. Please install antiword or catdoc, "
            f"or convert the file to .docx format. File: {file_path}"
        )
    
    def detect_language(self, text: str) -> str:
        """
        Detect the language of the text.
        
        Returns:
            Language code: 'uz-latn', 'uz-cyrl', 'ru', or 'mixed'
        """
        # Cyrillic characters
        cyrillic = set('абвгдеёжзийклмнопрстуфхцчшщъыьэюяАБВГДЕЁЖЗИЙКЛМНОПРСТУФХЦЧШЩЪЫЬЭЮЯ')
        # Latin characters
        latin = set('abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ')
        # Uzbek-specific Latin
        uzbek_latin = set("o'g'OʻGʻ")
        # Russian-specific (not in Uzbek Cyrillic)
        russian_specific = set('ъыэё')
        
        cyrillic_count = sum(1 for c in text if c in cyrillic)
        latin_count = sum(1 for c in text if c in latin)
        uzbek_latin_count = sum(1 for c in text if c in uzbek_latin)
        russian_count = sum(1 for c in text if c in russian_specific)
        
        total = cyrillic_count + latin_count
        
        if total == 0:
            return 'uz-latn'  # Default
        
        # Determine primary script
        if cyrillic_count > latin_count * 2:
            # Primarily Cyrillic
            if russian_count > 0:
                return 'ru'
            return 'uz-cyrl'
        elif latin_count > cyrillic_count * 2:
            # Primarily Latin
            return 'uz-latn'
        else:
            return 'mixed'


class PaddleOCRProcessor:
    """
    Alternative OCR processor using PaddleOCR.
    Better for complex layouts and Uzbek text.
    """
    
    def __init__(self):
        """Initialize PaddleOCR."""
        try:
            from paddleocr import PaddleOCR
            self.ocr = PaddleOCR(
                use_angle_cls=True,
                lang='cyrillic',  # Supports Cyrillic-based languages
                use_gpu=False
            )
            self.available = True
        except ImportError:
            logger.warning("PaddleOCR not installed")
            self.available = False
    
    def extract_text(self, image_path: str) -> Tuple[str, float]:
        """Extract text using PaddleOCR."""
        if not self.available:
            raise RuntimeError("PaddleOCR is not available")
        
        result = self.ocr.ocr(image_path)
        
        lines = []
        confidences = []
        
        for line in result[0]:
            text = line[1][0]
            confidence = line[1][1]
            lines.append(text)
            confidences.append(confidence)
        
        avg_confidence = sum(confidences) / len(confidences) if confidences else 0
        return "\n".join(lines), avg_confidence
